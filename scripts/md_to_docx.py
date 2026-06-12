"""Convert docs/API_COLLECTIONS.md to a Word .docx file.

Lightweight ad-hoc converter for this specific document. Handles the
markdown features actually used in API_COLLECTIONS.md: ATX headings,
GFM tables, fenced code blocks, blockquotes, ordered/unordered lists,
inline code, bold, links, and horizontal rules.
"""
from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SRC  = os.path.join(ROOT, "docs", "API_COLLECTIONS.md")
DST  = os.path.join(ROOT, "docs", "API_COLLECTIONS.docx")
LOGO = os.path.join(ROOT, "static", "logo_dark.png")
if not os.path.exists(LOGO):
    LOGO = os.path.join(ROOT, "static", "logo.png")

BLACK = RGBColor(0x00, 0x00, 0x00)


def force_black(run):
    """Force the run's font color to black, regardless of theme/style."""
    run.font.color.rgb = BLACK
    rPr = run._element.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), "000000")
    color.set(qn("w:themeColor"), "none")


# ── Inline parsing ──────────────────────────────────────────────────────────

INLINE_RE = re.compile(
    r"(`[^`]+`)"               # inline code
    r"|(\*\*[^*]+\*\*)"        # bold
    r"|(\[[^\]]+\]\([^)]+\))"  # markdown link
)


def add_inline(paragraph, text):
    """Render a line of inline-formatted text into the given paragraph.
    All runs are forced to black, regardless of style/theme."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            force_black(run)
        token = m.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:val"), "clear")
            shade.set(qn("w:color"), "auto")
            shade.set(qn("w:fill"), "F2F2F2")
            run._element.get_or_add_rPr().append(shade)
            force_black(run)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            force_black(run)
        elif token.startswith("["):
            mlink = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            label = mlink.group(1) if mlink else token
            run = paragraph.add_run(label)
            run.underline = True
            force_black(run)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        force_black(run)


# ── Block builders ──────────────────────────────────────────────────────────

_HEADING_PT = {1: 20, 2: 15, 3: 12.5, 4: 11}
_HEADING_SPACE_BEFORE_PT = {1: 14, 2: 14, 3: 12, 4: 8}
_HEADING_SPACE_AFTER_PT  = {1: 6,  2: 4,  3: 3,  4: 2}


def add_heading(doc, text, level):
    """Render an ATX heading with explicit size/spacing/page-break rules.

    Endpoint sections (H3 starting with `GET `) get a page break before so each
    endpoint starts on a fresh page — keeps the schema table + JSON example
    visually together."""
    lvl = min(level, 4)
    h = doc.add_heading(level=lvl)
    h.text = ""
    pf = h.paragraph_format
    pf.space_before = Pt(_HEADING_SPACE_BEFORE_PT[lvl])
    pf.space_after  = Pt(_HEADING_SPACE_AFTER_PT[lvl])
    pf.keep_with_next = True
    pf.keep_together  = True
    # Endpoint headings in the source are written as `### \`GET /path\``,
    # so strip backticks before matching.
    if lvl == 3 and text.strip().lstrip("`").startswith("GET "):
        pf.page_break_before = True

    add_inline(h, text)
    size_pt = _HEADING_PT[lvl]
    for run in h.runs:
        run.font.size = Pt(size_pt)
        run.bold = True
        force_black(run)


def add_paragraph(doc, text, *, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    add_inline(p, text)
    return p


def add_blockquote(doc, lines):
    """Render a contiguous run of '>' lines as a single quoted paragraph."""
    text = " ".join(line.lstrip("> ").rstrip() for line in lines if line.strip("> ").strip())
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "999999")
    pBdr.append(left)
    pPr.append(pBdr)
    add_inline(p, text)
    for run in p.runs:
        run.italic = True
        force_black(run)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent  = Cm(0.4)
    pf.right_indent = Cm(0.0)
    pf.space_before = Pt(4)
    pf.space_after  = Pt(8)
    pf.line_spacing = 1.0  # tight: each \n is exactly one line
    pf.keep_together = True  # don't split the JSON across pages mid-block
    pPr = p._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:color"), "auto")
    shade.set(qn("w:fill"), "F5F5F5")
    pPr.append(shade)
    run = p.add_run(code_text.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    force_black(run)


def add_list_item(doc, text, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(2)
    pf.left_indent  = Cm(0.6)
    pf.line_spacing = 1.15
    add_inline(p, text)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_table_row(line):
    """Split a `| a | b | c |` row into [a, b, c]."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _set_cell_borders(cell, color="000000", sz=4):
    """Single black border on every side of the cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        borders.append(b)
    tcPr.append(borders)


def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


# Column-width templates for known schema-table shapes (relative weights).
_COL_WIDTHS = {
    ("Campo", "Tipo", "Nullable", "Origem", "Descrição"): (3.0, 2.0, 1.5, 2.0, 7.0),
    ("Campo", "Tipo", "Nullable", "Descrição"):           (3.5, 2.0, 1.5, 9.0),
    ("Campo", "Tipo", "Descrição"):                       (3.5, 2.5, 10.0),
    ("Campo", "Tipo", "Nullable"):                        (4.0, 3.0, 2.0),
    ("Campo", "Tipo"):                                    (5.0, 11.0),
    ("Param", "Tipo", "Notas"):                           (4.0, 2.5, 9.5),
    ("HTTP", "code"):                                     (1.5, 14.5),
    ("Modo", "Param(s)", "Match"):                        (3.5, 4.5, 8.0),
}


def _strip_md(s: str) -> str:
    """Strip markdown decorations (backticks, asterisks) from a header cell so
    the column-width template lookup matches independently of formatting."""
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    return s.strip()


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Disable autofit at the XML level too — Word ignores the python-docx
    # setting in some cases.
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    # Column widths — match against header text with markdown stripped.
    header_key = tuple(_strip_md(h) for h in header)
    weights = _COL_WIDTHS.get(header_key)
    if weights is None:
        weights = tuple([1.0] * len(header))
    total_cm = 16.0
    weight_sum = sum(weights)
    widths = [Cm(total_cm * w / weight_sum) for w in weights]
    for j, w in enumerate(widths):
        for row in table.rows:
            row.cells[j].width = w

    # Header row.
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_borders(cell)
        _shade_cell(cell, "EAEAEA")
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after  = Pt(1)
        pf.line_spacing = 1.1
        add_inline(p, cell_text)
        for run in p.runs:
            run.bold = True
            force_black(run)

    # Body rows.
    for i, row in enumerate(rows, start=1):
        for j in range(len(header)):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after  = Pt(1)
            pf.line_spacing = 1.1
            value = row[j] if j < len(row) else ""
            add_inline(p, value)


# ── Document driver ─────────────────────────────────────────────────────────

def add_logo_header(doc, logo_path):
    """Embed the Beehus logo in the document's primary header.
    Aligned right; sized to ~3.8 cm wide so the rest of the header is empty."""
    if not os.path.exists(logo_path):
        return
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(logo_path, width=Cm(3.8))


def force_default_styles_black(doc):
    """Set every built-in style's color to black so the rendered output
    doesn't pick up the theme's blue heading color."""
    for style in doc.styles:
        try:
            font = style.font
        except AttributeError:
            continue
        try:
            font.color.rgb = BLACK
        except (AttributeError, ValueError):
            pass


def _setup_page(doc):
    """A4 portrait, balanced margins, room for the logo header."""
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width  = Mm(210)
    sec.left_margin  = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin    = Cm(3.0)  # leaves room for the logo header
    sec.bottom_margin = Cm(2.0)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)


def _tune_base_styles(doc):
    """Make the built-in heading and Normal styles match the per-paragraph
    overrides — so anyone editing the .docx sees consistent metrics."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after  = Pt(4)
    npf.line_spacing = 1.2

    for lvl in (1, 2, 3, 4):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Calibri"
        h.font.size = Pt(_HEADING_PT[lvl])
        h.font.bold = True
        h.font.color.rgb = BLACK
        pf = h.paragraph_format
        pf.space_before = Pt(_HEADING_SPACE_BEFORE_PT[lvl])
        pf.space_after  = Pt(_HEADING_SPACE_AFTER_PT[lvl])
        pf.keep_with_next = True


def convert(src_path, dst_path):
    with open(src_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    _setup_page(doc)
    _tune_base_styles(doc)
    force_default_styles_black(doc)
    add_logo_header(doc, LOGO)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip("\n")

        # Skip blank lines.
        if not line.strip():
            i += 1
            continue

        # Horizontal rule.
        if re.fullmatch(r"-{3,}", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Fenced code block.
        if line.lstrip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, "".join(buf))
            continue

        # ATX heading.
        m = re.match(r"^(#{1,6})\s+(.*?)\s*#*$", line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), level)
            i += 1
            continue

        # GFM table — header line followed by --- separator.
        if line.lstrip().startswith("|") and i + 1 < n and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            header = parse_table_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(parse_table_row(lines[i].rstrip("\n")))
                i += 1
            add_table(doc, header, rows)
            continue

        # Blockquote (collect contiguous '>' lines).
        if line.lstrip().startswith(">"):
            buf = [line]
            i += 1
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(lines[i].rstrip("\n"))
                i += 1
            add_blockquote(doc, buf)
            continue

        # Ordered list.
        m = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m:
            add_list_item(doc, m.group(2), ordered=True)
            i += 1
            continue

        # Unordered list.
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            add_list_item(doc, m.group(1), ordered=False)
            i += 1
            continue

        # Plain paragraph (collapse soft-wrapped lines into one).
        buf = [line]
        i += 1
        while i < n:
            nxt = lines[i].rstrip("\n")
            if not nxt.strip():
                break
            if (nxt.lstrip().startswith(("#", ">", "|", "```", "- ", "* "))
                    or re.match(r"^\s*\d+\.\s+", nxt)
                    or re.fullmatch(r"-{3,}", nxt.strip())):
                break
            buf.append(nxt)
            i += 1
        add_paragraph(doc, " ".join(buf))

    doc.save(dst_path)
    print(f"wrote {dst_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else SRC
    dst = sys.argv[2] if len(sys.argv) > 2 else DST
    convert(src, dst)
